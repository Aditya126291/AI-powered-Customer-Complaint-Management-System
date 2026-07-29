"""Build the interview-ready AIVOA video-script guide as a DOCX document."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "output" / "AIVOA_Complaint_Copilot_Integrated_Code_Explanation_Scripts.docx"

NAVY = "172554"
INDIGO = "4338CA"
SLATE = "475569"
PALE = "EEF2FF"
LIGHT = "F8FAFC"
ORANGE = "C2410C"


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    elem = OxmlElement("w:tblHeader")
    elem.set(qn("w:val"), "true")
    tr_pr.append(elem)


def keep_with_next(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    elem = OxmlElement("w:keepNext")
    p_pr.append(elem)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("AIVOA Complaint Copilot  |  ")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string(SLATE)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    paragraph._p.append(field)


def set_style(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string("1F2937")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.13

    for name, size, color in (("Title", 29, NAVY), ("Heading 1", 18, NAVY), ("Heading 2", 13, INDIGO), ("Heading 3", 11, ORANGE)):
        style = styles[name]
        style.font.name = "Aptos Display" if name != "Heading 3" else "Aptos"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), style.font.name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(13 if name != "Title" else 0)
        style.paragraph_format.space_after = Pt(6)


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    add_page_number(p)


def add_cover(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, NAVY)
    set_cell_margins(cell, top=420, start=420, bottom=420, end=420)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AIVOA\nCOMPLAINT COPILOT")
    r.font.name = "Aptos Display"
    r.font.size = Pt(29)
    r.font.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Two interview-ready video scripts")
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor.from_string("C7D2FE")

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI-Powered Customer Complaint Management System")
    r.font.size = Pt(16)
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Prepared from the current React, FastAPI, LangGraph, Groq, document-processing, and persistence implementation")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor.from_string(SLATE)

    doc.add_paragraph()
    box = doc.add_table(rows=1, cols=1)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = box.cell(0, 0)
    shade(cell, PALE)
    set_cell_margins(cell, top=160, start=180, bottom=160, end=180)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Use this guide as a spoken script. Items marked ON SCREEN are actions to perform while recording; the remaining text is ready to read aloud.")
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor.from_string("312E81")

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("Recommended recording length: ").bold = True
    meta.add_run("7–8 minutes per video\n")
    meta.add_run("Keep the application running at http://localhost:5173 before recording.")
    for run in meta.runs:
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string(SLATE)


def add_callout(doc, heading, text, fill=PALE, color="312E81"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, top=120, start=150, bottom=120, end=150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    h = p.add_run(heading.upper() + "  ")
    h.font.bold = True
    h.font.size = Pt(8.5)
    h.font.color.rgb = RGBColor.from_string(color)
    r = p.add_run(text)
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string(color)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_segment(doc, timestamp, on_screen, narration):
    heading = doc.add_paragraph(style="Heading 2")
    heading.add_run(timestamp)
    keep_with_next(heading)
    screen = doc.add_paragraph()
    screen.paragraph_format.space_after = Pt(4)
    label = screen.add_run("ON SCREEN  ")
    label.bold = True
    label.font.size = Pt(8.5)
    label.font.color.rgb = RGBColor.from_string(ORANGE)
    body = screen.add_run(on_screen)
    body.italic = True
    body.font.color.rgb = RGBColor.from_string(SLATE)
    say = doc.add_paragraph()
    say.paragraph_format.space_after = Pt(9)
    label = say.add_run("SAY  ")
    label.bold = True
    label.font.size = Pt(8.5)
    label.font.color.rgb = RGBColor.from_string(INDIGO)
    say.add_run(narration)


def add_code_cutaway(doc, file_and_target, on_screen, narration):
    """Add a brief code cutaway immediately after the related workflow explanation."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, "ECFDF5")
    set_cell_margins(cell, top=120, start=160, bottom=120, end=160)
    p = cell.paragraphs[0]
    r = p.add_run("CODE CUTAWAY  ")
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string("047857")
    r = p.add_run(file_and_target)
    r.bold = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string("065F46")
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("SHOW  ")
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string("047857")
    r = p.add_run(on_screen)
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string("065F46")
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("SAY  ")
    r.bold = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor.from_string("047857")
    r = p.add_run(narration)
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor.from_string("065F46")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_prompt_box(doc, title, prompts):
    doc.add_paragraph(title, style="Heading 2")
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["Moment", "Exact prompt or action"]
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell, NAVY)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        r = p.add_run(header)
        r.font.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
        r.font.size = Pt(9)
    set_repeat_table_header(table.rows[0])
    for moment, prompt in prompts:
        cells = table.add_row().cells
        for cell in cells:
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
        cells[0].paragraphs[0].add_run(moment).bold = True
        cells[0].paragraphs[0].runs[0].font.size = Pt(9)
        cells[1].paragraphs[0].add_run(prompt).font.size = Pt(9)
    doc.add_paragraph()


def add_workflow_video(doc):
    doc.add_page_break()
    doc.add_heading("Video 1 — Technical Workflow & Architecture", level=1)
    p = doc.add_paragraph()
    p.add_run("Target length: ").bold = True
    p.add_run("about 7 minutes. ")
    p.add_run("Goal: explain the problem, the end-to-end data flow, and the engineering decisions without sounding like you are reading a code walkthrough.")
    add_callout(doc, "Recording setup", "Start with the app visible. Keep a second tab ready for the FastAPI /docs page only if you want to show the API surface; it is optional.")

    add_segment(doc, "0:00–0:40 | Opening: the problem", "Show the two-panel AIVOA Complaint Copilot screen. Keep the empty complaint form and welcome message visible.", "Hello, and welcome to my project, AIVOA Complaint Copilot. It is an AI-powered customer complaint management system designed for API and finished dosage form quality-assurance teams. The problem it addresses is simple but important: complaints often arrive as unstructured emails, call notes, or documents, while a QMS needs a consistent, complete record. Instead of asking a user to manually copy every detail into a long form, this application lets them describe the complaint naturally or upload a document. The copilot extracts the important information, fills the correct fields, highlights what changed, and keeps the user in control of the final record.")

    add_segment(doc, "0:40–1:25 | User journey", "Point first to the Copilot panel, then sweep across the form sections on the left.", "The user journey has four stages. First, the user supplies a complaint by chat or file upload. Second, the backend extracts a controlled patch of form values rather than returning unstructured text. Third, the React interface applies that patch to the visible QMS form, so the user can immediately review or edit every field. Finally, once the record looks correct, the user commits it to the complaints ledger. This makes the AI a structured intake assistant, not an uncontrolled system of record. The human reviewer remains responsible for verification and final QA decisions.")

    add_segment(doc, "1:25–2:15 | Frontend design and state", "Show the left form and right chat. Briefly point out the purple highlighted inputs after a sample has been processed, if available.", "On the frontend, the application uses React with Redux Toolkit. The layout intentionally separates the controlled complaint form from the conversational copilot. The complaint slice holds the form values, changed-field indicators, save status, and the active database record identifier. The copilot slice holds the chat history, processing state, missing-field feedback, risk assessment, root-cause hypothesis, and CAPA suggestions. When the API sends back a patch, Redux applies only non-empty fields, and visually marks updated text inputs. That design makes it clear what the assistant has changed and allows a user to override any value manually.")
    add_code_cutaway(doc, "frontend/src/features/complaint/complaintSlice.ts — applyPatch()", "Switch to the code editor. Show applyPatch(), then highlight changedFields and the guarded check that ignores empty updates. Keep it on screen for 8–10 seconds.", "This is the Redux reducer that applies the AI result. Instead of directly changing inputs, I update a controlled state object and keep track of changed fields. That lets the UI highlight the exact values populated by the assistant, while still allowing a QA user to edit any value manually.")

    add_segment(doc, "2:15–3:10 | API boundary", "Optionally open http://localhost:8000/docs briefly, then return to the app. If not using Swagger, leave the app visible.", "The frontend communicates through a Vite development proxy to a FastAPI backend. The core conversational endpoint is POST slash API slash copilot slash process. It receives the latest message and the current form, then returns a structured response containing a message, a field patch, missing required fields, an initial risk assessment, a root-cause hypothesis, and CAPA recommendations. The upload endpoint receives multipart form data, reads the uploaded document, and sends its extracted text through the same intake pipeline. Separate endpoints commit, update, list, and delete complaint records. Keeping these responsibilities separate makes the system easier to test and extend.")
    add_code_cutaway(doc, "backend/app/main.py — CopilotRequest, CopilotResponse, and process_copilot()", "Open main.py and show the request/response models followed by the process_copilot endpoint. Do not scroll through the whole file.", "Here I define the API contract. The request includes both the new message and the current form context. The response returns a typed patch plus AI-assisted QA information. This is the boundary between the React application and the LangGraph workflow, so it keeps both sides predictable and easier to test.")

    add_segment(doc, "3:10–4:15 | LangGraph intake pipeline", "Show the project architecture diagram if you create one, or simply return to the app while explaining the node names on screen as captions.", "At the center of the backend is a LangGraph state workflow. The state begins with the user text and the current form. The first node, extract fields, calls Groq through LangChain with a Pydantic structured-output schema. The schema limits the model to approved QMS field names and captures the intent, updates, risk, root-cause hypothesis, and CAPA recommendations. The next node, normalize patch, standardizes dropdown-style values such as severity and priority. Then validate completeness merges the patch with the current form and reports missing essentials such as customer, product, batch or lot number, and defect summary. Finally, build message creates the helpful conversational response shown in the UI. The flow is extract fields, normalize patch, validate completeness, and build message.")
    add_code_cutaway(doc, "backend/app/agent.py — StructuredExtraction and build_intake_graph()", "Open agent.py. First show StructuredExtraction at the top; then jump directly to build_intake_graph() and pause on the graph.add_node and graph.add_edge lines.", "This is the core AI implementation. StructuredExtraction tells the model the exact response shape I expect, and build_intake_graph wires the processing into four explicit nodes. Showing this is important because it proves that the application uses a stateful, structured workflow instead of a single unstructured chatbot prompt.")

    add_segment(doc, "4:15–5:10 | AI reliability and safe fallbacks", "Keep the form visible; point to the ‘AI output must be verified’ text at the bottom of the Copilot panel.", "The Groq model is configured through environment variables, so the model choice and API key are not hardcoded in the frontend. To improve reliability, the backend combines model extraction with deterministic, high-precision pattern extraction for common complaint data such as batch numbers, product strengths, dates, quantities, and document labels. It also synthesizes a formal defect summary when the incoming message is too informal. If the AI provider is unavailable, the application falls back to local structured extraction plus form-aware root-cause and CAPA guidance, instead of failing the entire intake. This is why the UI explicitly states that AI output must be verified: the system assists QA triage; it does not replace QA judgment.")
    add_code_cutaway(doc, "backend/app/agent.py — extract_fields(), _deterministic_extract(), and the fallback return", "Stay in agent.py. Show the ChatGroq structured-output invocation, then the deterministic extraction merge directly below it. Finally point to the fallback return near the end of extract_fields().", "The important engineering decision here is resilience. I use Groq for intent and richer interpretation, but I also preserve deterministic extraction for identifiers such as batches and dates. If the external model call fails, this function still returns a usable patch and QA guidance rather than breaking the intake flow.")

    add_segment(doc, "5:10–5:55 | Document intake", "Hover over or click the upload zone without uploading yet.", "For document-driven complaints, the upload zone accepts text-based PDF, DOCX, TXT, and EML files up to ten megabytes. The backend uses file-type-specific parsers, extracts readable text, caps the extracted content at twenty thousand characters, and then passes that text to the same AI workflow. This avoids maintaining two separate extraction paths. One important limitation is that scanned or image-only PDFs need OCR, and OCR is deliberately not enabled in the current build. In an interview, I would call that out as a planned enhancement rather than misrepresenting the present capability.")
    add_code_cutaway(doc, "backend/app/main.py — extract_uploaded_document() and process_uploaded_document()", "Open main.py and show SUPPORTED_UPLOAD_EXTENSIONS, the ten-megabyte limit, and the file-format branches in extract_uploaded_document().", "This code proves that the upload feature is implemented on the backend. I validate the extension and file size, use the correct parser for PDF, DOCX, text, or email, reject unreadable image-only PDFs clearly, and then reuse the same intake workflow that powers chat messages.")

    add_segment(doc, "5:55–6:40 | Persistence, duplicate protection, and ledger", "Open Saved Complaints if records are present, then close it.", "When the reviewer clicks Save complaint, the backend creates a QMS record through SQLAlchemy. The database layer is configured for MySQL or PostgreSQL, with a SQLite development fallback so local demos continue working when a server database is not available. Each new record receives a generated complaint number in the format CC-year-sequence. Before committing, the backend checks matching customer, product, and batch information to block duplicate complaints. Saved records can be viewed in the ledger, opened for editing, and deleted. After a saved record is edited in the form, updates are automatically synchronized to the existing database record after a short debounce, rather than creating a second complaint.")
    add_code_cutaway(doc, "backend/app/main.py — commit_complaint(); frontend/src/components/ComplaintForm.tsx — auto-sync effect", "Show the duplicate-detection block and complaint-number generation in commit_complaint(). Then show the short useEffect with the 500-millisecond debounce in ComplaintForm.tsx.", "This is where I treated the system as a record-management product, not only a chatbot. Before I save, I check for the same customer, product, and batch. I generate a traceable complaint number, and after a saved record is edited I synchronize updates to the same record rather than creating another database entry.")

    add_segment(doc, "6:40–7:20 | Close and testing", "Return to the main screen. Optionally show the tests or README in a code editor after the last sentence.", "To validate the system, the project includes checks for backend health and Groq configuration, natural-language entity extraction, conversational updates, document upload extraction, frontend availability, and database persistence. The main value of this project is the complete workflow: a complaint begins as human language or a document, becomes a reviewable structured form, receives initial AI-assisted QA insights, and is stored as a traceable complaint record. That is the AIVOA Complaint Copilot workflow. Thank you for watching.")
    add_code_cutaway(doc, "test_final_suite.py and test_db_persistence.py", "Open the tests only long enough to show their named sections: health, natural-language intake, conversational update, upload extraction, frontend availability, and database persistence.", "I also wrote focused tests around the most important paths. This demonstrates that I verified the system beyond the happy-path user interface, including the AI connection, partial updates, uploads, and persistence behaviour.")


def add_demo_video(doc):
    doc.add_page_break()
    doc.add_heading("Video 2 — End-to-End Product Demonstration", level=1)
    p = doc.add_paragraph()
    p.add_run("Target length: ").bold = True
    p.add_run("about 8 minutes. ")
    p.add_run("Goal: demonstrate every major feature in a calm, deliberate sequence. Pause briefly after each AI response so viewers can see the form changes.")
    add_callout(doc, "Before you record", "Use a fresh complaint or click ‘File another complaint’ first. Keep the sample file ready: sample_documents/Apollo_Pharmacy_Discolored_Capsules.txt. If a duplicate warning appears because the same sample was saved earlier, use a different batch number or demonstrate the warning as a feature.", fill="FFF7ED", color="9A3412")

    add_segment(doc, "0:00–0:40 | Introduce the screen", "Show the app at localhost:5173. Point out the status pill, Saved Complaints button, form groups, upload box, chat history, and composer.", "This is the AIVOA Complaint Copilot interface. On the left is the controlled Log Customer Complaint form for API and finished dosage form quality assurance. It contains origin and customer details, product and batch identification, facility and material impact, complaint details, and initial severity and priority. On the right is the copilot. I can either type a natural-language complaint, like an email summary, or upload a supporting complaint document. The system then fills the form while preserving the ability to review or correct every value.")

    add_segment(doc, "0:40–1:55 | Natural-language intake", "Paste the initial intake prompt from the prompt table below into the chat box and press Send. Wait for the response, then slowly point through the filled fields.", "I will start with a typical customer report. Notice that I am not filling the form field by field. I am simply describing the issue in normal language. After I send this, the copilot processes the message and applies a structured patch to the form. We can now see the customer, source, product, strength, batch number, manufacturing and expiry dates, quantity, complaint type, severity, priority, and a formal defect summary. The highlighted input fields show what was updated. This is useful for speed, but the reviewer can still inspect the values before saving. If the assistant says a key field is missing, that feedback is also shown in the conversation.")

    add_segment(doc, "1:55–2:55 | Conversational correction", "Send the correction prompt. When the response returns, point only to batch/quantity and explain that the rest of the form stays available for review.", "Now I will demonstrate an important workflow: correcting a complaint after the initial intake. I do not need to start again or manually search through the form. I can send a short follow-up message saying that the batch number was entered incorrectly and specify the affected quantity. The backend receives the current form along with this message and returns a partial update. The corrected fields change here, while the rest of the complaint context remains in the form. This supports the way people actually work with customer reports, where details are often clarified in multiple messages.")

    add_segment(doc, "2:55–3:45 | Manual review and controlled values", "Click the Severity and Priority dropdowns briefly. Do not necessarily alter them; show their available options. Optionally type directly into Structured Defect Summary to demonstrate human override.", "The assistant speeds up intake, but this remains a controlled QA form. Severity and priority use predefined dropdown options, including Low, Medium, High, Critical or Urgent, and Needs QA Review. A user can adjust them directly after considering the evidence. The structured defect summary is editable too, so the QA reviewer can refine the language before committing the record. This review step is intentional: AI suggestions are useful initial triage, but the final classification requires human oversight.")

    add_segment(doc, "3:45–4:55 | Document upload", "Click the upload area and select sample_documents/Apollo_Pharmacy_Discolored_Capsules.txt. Wait for processing and show the chat message naming the uploaded file and the populated fields.", "Next, I will demonstrate document intake. The upload zone accepts text-based PDF, Word, text, and email files. I am uploading the provided Apollo Pharmacy complaint text file. The backend validates the file type and size, extracts readable text, and runs that extracted content through the same intake workflow. The form is populated from the document, and the assistant also provides the initial risk and QA recommendations. This is valuable when complaints arrive as attachments rather than being copied into a chat. For transparency, scanned image-only PDFs are not currently supported because OCR is not enabled in this version.")

    add_segment(doc, "4:55–6:10 | QA quick actions", "With a completed complaint on screen, click Summarize complaint, AI risk classification, Root cause recommendation, and CAPA recommendation one at a time. Give viewers a second to read each chat response.", "Once the complaint has enough detail, the quick-action controls appear below the chat. The first gives an executive summary using the customer, product, batch, quantity, incident summary, severity, and priority already in the form. The risk classification surfaces the initial triage assessment. The root-cause recommendation presents a scientific hypothesis and a suggested QA investigation plan. Finally, the CAPA recommendation provides practical containment, root-cause investigation, and preventive-action steps. These are not final decisions or automatic closures. They are structured starting points for a quality team to review, investigate, and document under its own QMS procedures.")

    add_segment(doc, "6:10–7:00 | Save and duplicate protection", "Click Save complaint. Show the success message, record number, status pill, and Saved Complaints count. If a duplicate warning appears, explain it and use a new batch for the next attempt.", "Now that the record has been reviewed, I will save it. The application commits the form and its current risk assessment to the database and assigns a traceable complaint number. The status changes to Saved, and the assistant confirms the number in the conversation. The backend also checks for duplicate complaints using the key customer, product, and batch combination. If the same complaint is submitted again, it is blocked rather than silently creating two QMS records. That is a small but important safeguard for data integrity.")

    add_segment(doc, "7:00–7:45 | Saved ledger and record editing", "Click Saved Complaints, select the saved record, and point out the structured detail table. Click Edit Record, close the modal, change one safe value such as the complaint source, wait a moment, then reopen the ledger to verify it remains the same record.", "The Saved Complaints ledger gives the user a compact view of the stored database records. Each record shows its number, severity, product, batch, and customer, while the details panel shows the complete saved fields, risk assessment, and committed status. I can select Edit Record to load a saved complaint back into the form. After that, changes are synchronized to the existing record rather than creating a new one. The ledger also provides a delete action for removing a record when appropriate in this prototype.")

    add_segment(doc, "7:45–8:15 | Reset and closing", "Click File another complaint if it is visible; otherwise explain the Reset form button. Finish on the clean interface.", "To begin a completely new intake, I can use File another complaint after saving, or Reset form while still drafting. That clears the form and the copilot context for the next customer issue. In summary, AIVOA Complaint Copilot combines conversational intake, document extraction, controlled form review, AI-assisted QA insights, and persistent complaint records in one workflow. It reduces repetitive data entry while keeping the human reviewer at the center of quality decisions. Thank you for watching the demonstration.")

    add_prompt_box(doc, "Exact prompts and recording actions", [
        ("Initial chat intake", "Apollo Pharmacy reported discolored capsules in Amoxicillin Capsules 500 mg. Batch number AMX240602. Manufacturing date March 2026. Expiry date February 2028. The complaint came from a retail pharmacy outlet. Affected quantity is 48 capsules. Please log this complaint."),
        ("Follow-up correction", "Correction: the batch number is BMX240602 and the affected quantity is 48 capsules."),
        ("Optional detail request", "Please add that the observed issue is brown discoloration in sealed customer boxes and set the complaint source to retail pharmacy outlet."),
        ("Document demo", "Upload the file at sample_documents/Apollo_Pharmacy_Discolored_Capsules.txt through the drag-and-drop area."),
        ("Quick actions", "Click each pill in sequence: Summarize complaint, AI risk classification, Root cause recommendation, and CAPA recommendation."),
        ("Persistence demo", "Click Save complaint, then Saved Complaints. Select the record and use Edit Record to demonstrate loading and synchronizing an existing complaint."),
    ])


def add_final_notes(doc):
    doc.add_heading("Interview delivery notes", level=1)
    add_callout(doc, "One-line positioning", "AIVOA Complaint Copilot converts conversational or document-based pharmaceutical complaints into reviewable, traceable QMS records while keeping QA staff in control.", fill="E0E7FF")
    bullets = [
        "Speak at a measured pace and pause after each AI response. The target timings include those pauses.",
        "Never read API keys, database URLs, or environment variables aloud or show them in a recording.",
        "Say ‘initial recommendation’ for risk, root cause, and CAPA outputs. Do not call them final QA decisions.",
        "For uploads, say ‘text extraction’ rather than OCR. The current version explicitly does not OCR image-only PDFs.",
        "If the same sample is already in the database, the duplicate block is expected behaviour. Use it as evidence of the data-integrity control or change the batch number for a clean save.",
    ]
    for item in bullets:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_code_tour(doc):
    doc.add_page_break()
    doc.add_heading("Code Tour Appendix — What to Show the Interviewer", level=1)
    p = doc.add_paragraph()
    p.add_run("Use this as a 90-second code walkthrough after the architecture explanation. ").bold = True
    p.add_run("Do not scroll through every file. Open three to five focused locations, explain the engineering decision in each, then return to the running application.")
    add_callout(doc, "Best sequence", "Start with the LangGraph agent, move to the FastAPI upload and persistence endpoints, then show one React/Redux state update. This tells a complete story from AI reasoning to user interface to database.")

    spotlights = [
        (
            "1. The controlled AI output contract",
            "Show: backend/app/agent.py — StructuredExtraction (near the top of the file).",
            "Why it matters: This Pydantic model defines the exact shape expected from the LLM: create or update intent, allowed form updates, a summary, risk, a root-cause hypothesis, CAPA recommendations, and dropdown intent. It demonstrates that the model is not being asked for free-form text only.",
            "Say: ‘I use a structured-output schema so the model returns machine-readable complaint data. The frontend receives a controlled patch rather than trying to parse an arbitrary AI paragraph.’",
        ),
        (
            "2. AI plus deterministic extraction",
            "Show: backend/app/agent.py — extract_fields(), especially the ChatGroq structured-output call and the deterministic merge immediately after it.",
            "Why it matters: The code combines Groq structured output with high-precision extraction for dates, batch numbers, quantities, and common document labels. This protects common QMS identifiers when a model response is incomplete.",
            "Say: ‘I did not rely on a single model response. I merge it with deterministic extraction for high-confidence entities, then synthesize a formal defect summary when the supplied text is too informal.’",
        ),
        (
            "3. The actual LangGraph workflow",
            "Show: backend/app/agent.py — build_intake_graph() and run_intake().",
            "Why it matters: The graph explicitly connects extract_fields, normalize_patch, validate_completeness, and build_message. This is a real stateful workflow, not a single hardcoded prompt.",
            "Say: ‘LangGraph makes the processing stages visible and extensible. Each node has one responsibility: extract, normalize, validate, and then communicate the result to the user.’",
        ),
        (
            "4. Document security and format handling",
            "Show: backend/app/main.py — SUPPORTED_UPLOAD_EXTENSIONS, extract_uploaded_document(), and process_uploaded_document().",
            "Why it matters: These functions restrict file types, enforce a ten-megabyte size limit, parse PDF, DOCX, TXT, and EML differently, cap extracted text, and give a clear message for image-only PDFs that need OCR.",
            "Say: ‘The upload feature is a backend workflow, not just a file picker. I validate files, extract text according to format, and pass that text through the same AI pipeline as chat input.’",
        ),
        (
            "5. Resilient API behaviour",
            "Show: backend/app/main.py — process_copilot().",
            "Why it matters: If the AI call throws an exception, the endpoint still returns a usable deterministic patch, initial risk, root cause, and CAPA guidance instead of giving the user a generic failure.",
            "Say: ‘I designed an explicit fallback path. The form can still receive structured information when the external AI service has a temporary problem, which makes the demo and the workflow more reliable.’",
        ),
        (
            "6. Data integrity at save time",
            "Show: backend/app/main.py — commit_complaint(), focusing on the duplicate check and CC-year-sequence number generation.",
            "Why it matters: It prevents the same customer/product/batch combination from being committed twice, creates a traceable complaint number, and stores risk alongside the form data.",
            "Say: ‘Saving is more than inserting a row. I added duplicate protection and traceable complaint numbering because these details matter in a QMS-style workflow.’",
        ),
        (
            "7. State update from an AI patch",
            "Show: frontend/src/components/Copilot.tsx — applyAiResult(), submit(), and upload(). Then show frontend/src/features/complaint/complaintSlice.ts — applyPatch().",
            "Why it matters: The UI sends current context, receives a typed patch, applies non-empty values to Redux state, and records which form fields changed. Chat and uploads share the same update route.",
            "Say: ‘The frontend does not directly manipulate random DOM elements. The assistant result goes through Redux, so the form remains a controlled stateful interface and changes are visible to the reviewer.’",
        ),
        (
            "8. Editing a committed record without duplication",
            "Show: frontend/src/components/ComplaintForm.tsx — the useEffect beginning with the auto-sync comment, plus handleSaveOrUpdate().",
            "Why it matters: Once a saved record is loaded for editing, the component debounces updates and calls the update endpoint rather than committing another complaint.",
            "Say: ‘I also handled the lifecycle after save. An edited record keeps its identity and automatically syncs to the database, instead of creating duplicate records.’",
        ),
        (
            "9. Portable persistence configuration",
            "Show: backend/app/database.py — DATABASE_URL selection and the SQLite fallback. Optionally show backend/app/models.py — ComplaintRecord.",
            "Why it matters: Production-style MySQL or PostgreSQL can be configured through the environment, while a SQLite fallback keeps local development and interviews runnable without a separate database server.",
            "Say: ‘The persistence layer is environment-driven. I can use MySQL or PostgreSQL in deployment, but local development still works predictably through SQLite.’",
        ),
        (
            "10. Evidence of testing",
            "Show: test_final_suite.py — the named test sections only; do not linger on every assertion.",
            "Why it matters: The suite covers health and Groq status, natural-language intake, conversational corrections, document extraction, and frontend availability. A separate test checks persistence.",
            "Say: ‘I added checks around the most important flows: the service configuration, extraction, update behaviour, file input, frontend availability, and database persistence.’",
        ),
    ]

    for title, show, why, say in spotlights:
        doc.add_heading(title, level=2)
        p = doc.add_paragraph()
        label = p.add_run("OPEN  ")
        label.bold = True
        label.font.color.rgb = RGBColor.from_string(ORANGE)
        p.add_run(show)
        p = doc.add_paragraph()
        label = p.add_run("WHAT IT PROVES  ")
        label.bold = True
        label.font.color.rgb = RGBColor.from_string(INDIGO)
        p.add_run(why)
        p = doc.add_paragraph()
        label = p.add_run("SAY  ")
        label.bold = True
        label.font.color.rgb = RGBColor.from_string(NAVY)
        p.add_run(say)

    doc.add_heading("Suggested 90-second code-show flow", level=2)
    flow = [
        "Open agent.py and show StructuredExtraction, then build_intake_graph. This establishes the AI workflow and its controlled response contract.",
        "Open main.py and show process_uploaded_document plus commit_complaint. This proves your API, validation, file handling, duplicate protection, and persistence work.",
        "Open Copilot.tsx and complaintSlice.ts. This proves the frontend receives structured responses and updates the controlled form through Redux.",
        "Finish in the running app by sending one message and saving one record. The code and product demonstration now reinforce each other.",
    ]
    for item in flow:
        doc.add_paragraph(item, style="List Number")


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.72)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    add_footer(section)
    set_style(doc)
    doc.core_properties.title = "AIVOA Complaint Copilot — Two Video Scripts"
    doc.core_properties.subject = "Interview-ready technical workflow and product demonstration scripts"
    doc.core_properties.author = "AIVOA Project Team"

    add_cover(doc)
    add_workflow_video(doc)
    add_demo_video(doc)
    add_final_notes(doc)
    add_code_tour(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
